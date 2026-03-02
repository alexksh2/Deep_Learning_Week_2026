"""
RAG pipeline CLI — PDF indexing and querying for the coaching chatbot.

Usage:
    # Build FAISS index from a PDF
    python3 rag_cli.py index --file resume.pdf --index-dir /tmp/rag_abc123

    # Query against an existing index
    python3 rag_cli.py query --question "What is the budget for AI?" --index-dir /tmp/rag_abc123

All diagnostic output goes to stderr. Only the final JSON result goes to stdout.

Dependencies:
    pip install -r ml-development/rag-pipeline/requirements.txt
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from functools import lru_cache
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Dict, List, Optional

from dotenv import load_dotenv
from openai import OpenAI

from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Load .env.local for standalone use
for _cand in [
    Path(__file__).parent.parent / ".env.local",
    Path(__file__).parent.parent / ".env",
]:
    if _cand.exists():
        load_dotenv(_cand)
        break

# ── Config ────────────────────────────────────────────────────────────────────

CHUNK_SIZE    = 500
CHUNK_OVERLAP = 100
TOP_K         = 5
EMBED_MODEL   = "all-MiniLM-L6-v2"

OPENAI_MODEL = os.getenv("OPENAI_MODEL_RAG", os.getenv("OPENAI_MODEL", "gpt-4o-mini"))

CITATION_PATTERN = re.compile(r"\[(c\d+\.\d+\.\d+(?:,\s*c\d+\.\d+\.\d+)*)\]")


# ── OpenAI helper ─────────────────────────────────────────────────────────────

@lru_cache(maxsize=1)
def get_openai_client() -> OpenAI:
    key = os.getenv("OPENAI_API_KEY")
    if not key:
        raise EnvironmentError("OPENAI_API_KEY not set")
    return OpenAI(api_key=key)


def call_openai(prompt: str, system: str = "", max_tokens: int = 1024) -> str:
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    resp = get_openai_client().chat.completions.create(
        model=OPENAI_MODEL,
        messages=messages,
        max_tokens=max_tokens,
        temperature=0.0,
    )
    return (resp.choices[0].message.content or "").strip()


# ── Phase 1 — Load & Chunk ────────────────────────────────────────────────────
# Extracted from dlw_workshop.py

def load_pdf(file_path: str) -> List[Document]:
    loader = PyPDFLoader(file_path)
    documents = loader.load()
    print(f"Loaded {len(documents)} pages from {file_path}", file=sys.stderr)
    return documents


def load_txt_or_docx(file_path: str) -> List[Document]:
    """Fallback loader for non-PDF files."""
    ext = Path(file_path).suffix.lower()
    if ext in (".docx", ".doc"):
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        with open(file_path, encoding="utf-8", errors="ignore") as f:
            text = f.read()
    return [Document(page_content=text, metadata={"page": 1})]


def load_document(file_path: str) -> List[Document]:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    return load_txt_or_docx(file_path)


def chunk_documents(
    documents: List[Document],
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> List[Document]:
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    chunks = text_splitter.split_documents(documents)
    print(f"Created {len(chunks)} chunks", file=sys.stderr)
    return chunks


# ── Phase 2 — Citation Metadata ───────────────────────────────────────────────
# Extracted from dlw_workshop.py

@dataclass
class BoundingBox:
    x0: float
    y0: float
    x1: float
    y1: float

    def to_dict(self) -> Dict:
        return asdict(self)


@dataclass
class Citation:
    citation_id:  str
    page_number:  int
    paragraph_id: str
    bbox:         Optional[BoundingBox] = None
    text_preview: str = ""

    def to_dict(self) -> Dict:
        return {
            "citation_id":  self.citation_id,
            "page_number":  self.page_number,
            "paragraph_id": self.paragraph_id,
            "bbox":         self.bbox.to_dict() if self.bbox else None,
            "text_preview": self.text_preview,
        }


class CitationExtractor:
    def __init__(self, pdf_path: Optional[str] = None):
        self.pdf_path = pdf_path
        self.pdf = None
        if pdf_path and PDFPLUMBER_AVAILABLE:
            try:
                self.pdf = pdfplumber.open(pdf_path)
            except Exception as e:
                print(f"Could not open PDF with pdfplumber: {e}", file=sys.stderr)

    def _generate_citation_id(self, page: int, para_idx: int) -> str:
        return f"c{page}.0.{para_idx}"

    def _extract_page_number(self, document: Document) -> int:
        if hasattr(document, "metadata") and document.metadata:
            return document.metadata.get("page", 0)
        return 0

    def _estimate_bbox(self, text_length: int, page: int) -> BoundingBox:
        if self.pdf and 0 < page <= len(self.pdf.pages):
            try:
                pdf_page = self.pdf.pages[page - 1]
                chars = pdf_page.chars
                if chars:
                    bbox_info = [{"x0": c["x0"], "y0": c["y0"], "x1": c["x1"], "y1": c["y1"]}
                                 for c in chars[:text_length]]
                    return BoundingBox(
                        x0=min(b["x0"] for b in bbox_info),
                        y0=min(b["y0"] for b in bbox_info),
                        x1=max(b["x1"] for b in bbox_info),
                        y1=max(b["y1"] for b in bbox_info),
                    )
            except Exception:
                pass

        avg_chars_per_line = 80
        line_height = 12
        lines = max(1, text_length // avg_chars_per_line)
        y_pos = 50 + ((page % 5) * 150)
        return BoundingBox(x0=50, y0=y_pos, x1=562, y1=y_pos + (lines * line_height))

    def enrich_chunk(self, document: Document, chunk_index: int) -> Document:
        page_num    = self._extract_page_number(document)
        citation_id = self._generate_citation_id(page_num, chunk_index)
        bbox        = self._estimate_bbox(len(document.page_content), page_num)
        citation    = Citation(
            citation_id=citation_id,
            page_number=page_num,
            paragraph_id=f"P{page_num}-Para{chunk_index}",
            bbox=bbox,
            text_preview=document.page_content[:100],
        )
        if not document.metadata:
            document.metadata = {}
        document.metadata.update({
            "citation_id":   citation_id,
            "page_number":   page_num,
            "paragraph_id":  citation.paragraph_id,
            "citation_bbox": bbox.to_dict(),
            "citation_data": citation.to_dict(),
        })
        return document


def chunk_documents_citation(
    documents: List[Document],
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    pdf_path: Optional[str] = None,
) -> List[Document]:
    chunks    = chunk_documents(documents, chunk_size, chunk_overlap)
    extractor = CitationExtractor(pdf_path)
    enriched  = [extractor.enrich_chunk(chunk, i) for i, chunk in enumerate(chunks)]
    print(f"Enriched {len(enriched)} chunks with citation metadata", file=sys.stderr)
    return enriched


# ── Vector store ──────────────────────────────────────────────────────────────

def create_vector_store(chunks: List[Document]) -> FAISS:
    print(f"Building FAISS vector store with {EMBED_MODEL}…", file=sys.stderr)
    embeddings   = HuggingFaceEmbeddings(model_name=EMBED_MODEL)
    vector_store = FAISS.from_documents(chunks, embeddings)
    print("Vector store ready", file=sys.stderr)
    return vector_store


def load_vector_store(index_dir: str) -> FAISS:
    print(f"Loading vector store from {index_dir}…", file=sys.stderr)
    embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)
    return FAISS.load_local(index_dir, embeddings, allow_dangerous_deserialization=True)


# ── Retrieval ─────────────────────────────────────────────────────────────────

def retrieval_with_citations(vector_store: FAISS, query: str, k: int = TOP_K) -> List[Document]:
    results   = vector_store.similarity_search(query, k=k)
    extractor = CitationExtractor()
    for idx, chunk in enumerate(results):
        if "citation_id" not in chunk.metadata:
            chunk = extractor.enrich_chunk(chunk, idx)
    print(f"Retrieved {len(results)} chunks", file=sys.stderr)
    return results


# ── Prompt ────────────────────────────────────────────────────────────────────
# Extracted from dlw_workshop.py format_citation_aware_prompt

SYSTEM_PROMPT = (
    "You are a helpful assistant that answers questions based on provided documents. "
    "Answer only based on the context provided. "
    "If the context does not contain the answer, say 'I don't know.'\n\n"
    "Use citations in your answer to reference the sources used to answer the question.\n\n"
    "IMPORTANT - Citation Guidelines:\n"
    "- When making a claim, immediately cite the source using its Citation ID in square brackets.\n"
    "- Place citations inline right after the relevant text: e.g. \"Budget will increase by 10% [c3.0.2].\"\n"
    "- If a claim uses multiple sources, list all: e.g. \"[c3.0.2, c4.0.1].\""
)


def format_prompt(question: str, chunks: List[Document]) -> str:
    context = "Retrieved documents with citations:\n"
    for i, chunk in enumerate(chunks, 1):
        citation_id  = chunk.metadata.get("citation_id", f"Doc{i}")
        page         = chunk.metadata.get("page_number", "?")
        paragraph_id = chunk.metadata.get("paragraph_id", "")
        context += f"\n{'=' * 60}\n"
        context += f"[Document {i}]\n"
        context += f"Citation ID: {citation_id}\n"
        context += f"Page: {page}\n"
        context += f"Paragraph ID: {paragraph_id}\n"
        context += f"{'=' * 60}\n"
        context += f"{chunk.page_content}\n"

    return f"{context}\n\nQuestion: {question}\n\nAnswer:"


# ── Citation linking ──────────────────────────────────────────────────────────
# Extracted from dlw_workshop.py CitationLinker

class CitationLinker:
    def link_citations(self, response: str, retrieved_chunks: List[Document]) -> Dict[str, Any]:
        citation_lookup: Dict[str, Dict] = {}
        for chunk in retrieved_chunks:
            cid = chunk.metadata.get("citation_id")
            if cid:
                citation_lookup[cid] = {
                    "page":         chunk.metadata.get("page_number"),
                    "paragraph_id": chunk.metadata.get("paragraph_id"),
                    "bbox":         chunk.metadata.get("citation_bbox"),
                    "preview":      chunk.page_content[:200],
                    "full_content": chunk.page_content,
                }

        cited_ids: List[str] = []
        for match in CITATION_PATTERN.findall(response):
            for cid in match.split(","):
                cid = cid.strip()
                if cid not in cited_ids:
                    cited_ids.append(cid)

        citations = [
            {"id": cid, "source": citation_lookup[cid]}
            for cid in cited_ids
            if cid in citation_lookup
        ]
        return {"response": response, "citations": citations, "sources": citation_lookup}


# ── Phase 3 — Claim-level LLM-as-a-Judge ─────────────────────────────────────
# Extracted from dlw_workshop.py, LLM calls adapted to OpenAI

def _extract_citation_ids(text: str) -> List[str]:
    ids: List[str] = []
    for match in CITATION_PATTERN.findall(text):
        for cid in match.split(","):
            cid = cid.strip()
            if cid not in ids:
                ids.append(cid)
    return ids


def extract_claims(response_text: str) -> List[Dict[str, Any]]:
    sentences = re.split(r"(?<=[.!?])\s+", response_text.strip())
    claims: List[Dict[str, Any]] = []
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        if len(sentence) < 10 and claims:
            claims[-1]["text"] += " " + sentence
            claims[-1]["citation_ids"] = _extract_citation_ids(claims[-1]["text"])
            claims[-1]["text"] = CITATION_PATTERN.sub("", claims[-1]["text"]).strip()
            continue
        citation_ids = _extract_citation_ids(sentence)
        clean_text   = CITATION_PATTERN.sub("", sentence).strip()
        clean_text   = re.sub(r"\s+([.!?,;:])", r"\1", clean_text)
        if clean_text:
            claims.append({"text": clean_text, "citation_ids": citation_ids})
    return claims


EVALUATION_SYSTEM_PROMPT = """You are a fact-checking assistant. Your job is to evaluate whether claims are supported by provided source texts.

For each claim you receive, output a JSON object with these exact fields:
- "claim_index": the integer index of the claim (as given)
- "support_level": one of exactly "supported", "partial", "unsupported", or "contradicted"
- "confidence": a float between 0.0 and 1.0
- "reasoning": one or two sentences explaining your decision

Output ONLY a valid JSON array of these objects. Do not include any explanation, markdown, or other text outside the JSON array."""


def build_evaluation_prompt(claims: List[Dict[str, Any]], citation_lookup: Dict[str, Dict]) -> str:
    prompt      = "Evaluate the following claims against their sources:\n"
    claim_index = 0
    for claim in claims:
        if not claim["citation_ids"]:
            continue
        claim_index += 1
        prompt += f"\n---\nClaim {claim_index}: {claim['text']}\n"
        for cid in claim["citation_ids"]:
            if cid in citation_lookup:
                content = citation_lookup[cid].get("full_content", citation_lookup[cid].get("preview", ""))
                if len(content) > 500:
                    content = content[:500] + "..."
                prompt += f"Source ({cid}): {content}\n"
    prompt += f"\nReturn a JSON array with {claim_index} objects, one per claim."
    return prompt


def _parse_evaluation_response(raw: str) -> List[Dict]:
    try:
        parsed = json.loads(raw.strip())
        if isinstance(parsed, list):
            return parsed
        if isinstance(parsed, dict):
            for v in parsed.values():
                if isinstance(v, list):
                    return v
    except (json.JSONDecodeError, ValueError):
        pass

    match = re.search(r"\[.*\]", raw, re.DOTALL)
    if match:
        try:
            parsed = json.loads(match.group())
            if isinstance(parsed, list):
                return parsed
        except (json.JSONDecodeError, ValueError):
            pass

    print("Could not parse evaluator response as JSON", file=sys.stderr)
    return []


def run_evaluation(
    response_text: str,
    citation_lookup: Dict[str, Dict],
) -> Optional[Dict[str, Any]]:
    claims = extract_claims(response_text)
    if not claims:
        return None

    cited_claims   = [c for c in claims if c["citation_ids"]]
    uncited_claims = [c for c in claims if not c["citation_ids"]]

    evaluations: List[Dict[str, Any]] = [
        {"text": c["text"], "citation_ids": [], "support_level": "uncited",
         "confidence": 0.0, "reasoning": "No citation provided for this claim."}
        for c in uncited_claims
    ]

    if cited_claims:
        user_prompt = build_evaluation_prompt(cited_claims, citation_lookup)
        try:
            raw       = call_openai(user_prompt, system=EVALUATION_SYSTEM_PROMPT, max_tokens=1000)
            raw       = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
            raw       = re.sub(r"\s*```\s*$", "", raw)
            eval_data = _parse_evaluation_response(raw)
        except Exception as e:
            print(f"Evaluation call failed: {e}", file=sys.stderr)
            eval_data = []

        for i, claim in enumerate(cited_claims):
            claim_idx = i + 1
            match = next((item for item in eval_data if item.get("claim_index") == claim_idx), None)
            if match is None and i < len(eval_data):
                match = eval_data[i]

            if match:
                support = match.get("support_level", "partial")
                if support not in ("supported", "partial", "unsupported", "contradicted"):
                    support = "partial"
                evaluations.append({
                    "text":          claim["text"],
                    "citation_ids":  claim["citation_ids"],
                    "support_level": support,
                    "confidence":    float(max(0.0, min(1.0, match.get("confidence", 0.5)))),
                    "reasoning":     match.get("reasoning", ""),
                })
            else:
                evaluations.append({
                    "text":          claim["text"],
                    "citation_ids":  claim["citation_ids"],
                    "support_level": "partial",
                    "confidence":    0.5,
                    "reasoning":     "Evaluator did not return a result for this claim.",
                })

    cited_evals = [e for e in evaluations if e["support_level"] != "uncited"]
    overall     = (
        sum(e["confidence"] for e in cited_evals) / len(cited_evals)
        if cited_evals else 0.0
    )

    return {
        "overall_confidence": round(overall, 2),
        "supported":   sum(1 for e in evaluations if e["support_level"] == "supported"),
        "partial":     sum(1 for e in evaluations if e["support_level"] == "partial"),
        "unsupported": sum(1 for e in evaluations if e["support_level"] == "unsupported"),
        "uncited":     sum(1 for e in evaluations if e["support_level"] == "uncited"),
        "claims":      evaluations,
    }


# ── Index action ──────────────────────────────────────────────────────────────

def action_index(file_path: str, index_dir: str) -> None:
    documents = load_document(file_path)
    if not documents:
        print(json.dumps({"error": "Could not extract text from file"}))
        sys.exit(1)

    chunks = chunk_documents_citation(documents, pdf_path=file_path)

    vector_store = create_vector_store(chunks)
    Path(index_dir).mkdir(parents=True, exist_ok=True)
    vector_store.save_local(index_dir)
    print(f"Index saved to: {index_dir}", file=sys.stderr)

    print(json.dumps({
        "status":    "ok",
        "chunks":    len(chunks),
        "pages":     len(documents),
        "index_dir": index_dir,
    }))


# ── Query action ──────────────────────────────────────────────────────────────

def action_query(question: str, index_dir: str, run_eval: bool) -> None:
    if not Path(index_dir).exists():
        print(json.dumps({"error": f"Index not found at: {index_dir}"}))
        sys.exit(1)

    vector_store = load_vector_store(index_dir)
    retrieved    = retrieval_with_citations(vector_store, question)

    prompt   = format_prompt(question, retrieved)
    print(f"Calling OpenAI ({OPENAI_MODEL})…", file=sys.stderr)
    response = call_openai(prompt, system=SYSTEM_PROMPT)

    linker = CitationLinker()
    linked = linker.link_citations(response, retrieved)

    evaluation = None
    if run_eval and linked["citations"]:
        print("Running claim-level evaluation…", file=sys.stderr)
        evaluation = run_evaluation(linked["response"], linked["sources"])

    # Flatten citations for the API (id + page + text preview)
    citations_out = [
        {"id": c["id"], "page": c["source"].get("page"), "text": c["source"].get("preview", "")}
        for c in linked["citations"]
    ]

    print(json.dumps({
        "answer":     linked["response"],
        "citations":  citations_out,
        "evaluation": evaluation,
    }))


# ── Entry point ───────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="RAG pipeline CLI")
    sub    = parser.add_subparsers(dest="action", required=True)

    idx_p = sub.add_parser("index", help="Build FAISS index from a document")
    idx_p.add_argument("--file",      required=True, help="Path to PDF / DOCX / TXT")
    idx_p.add_argument("--index-dir", required=True, help="Directory to save the index")

    qry_p = sub.add_parser("query", help="Answer a question against an existing index")
    qry_p.add_argument("--question",    required=True, help="Question to answer")
    qry_p.add_argument("--index-dir",   required=True, help="Directory containing the index")
    qry_p.add_argument("--no-evaluate", action="store_true",
                       help="Skip LLM-as-a-Judge evaluation (faster)")

    args = parser.parse_args()

    if args.action == "index":
        action_index(args.file, args.index_dir)
    elif args.action == "query":
        action_query(args.question, args.index_dir, not args.no_evaluate)


if __name__ == "__main__":
    main()
